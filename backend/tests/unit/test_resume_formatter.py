# File: backend/tests/unit/test_resume_formatter.py

import pytest
from app.services.resume_formatter import generate_formatted_resume_docx
from docx import Document
from io import BytesIO

def extract_text_from_docx(doc_bytes: BytesIO) -> list:
    doc = Document(doc_bytes)
    return [para.text for para in doc.paragraphs if para.text.strip()]

def test_generate_docx_valid_summary():
    text = """
    John Doe
    johndoe@email.com | 123-456-7890

    SUMMARY
    Experienced QA Engineer with automation skills.
    """
    result = generate_formatted_resume_docx(text, is_user_approved=False)
    assert isinstance(result, BytesIO)

    lines = extract_text_from_docx(result)
    assert "JOHN DOE" in lines[0]
    assert "SUMMARY" in lines
    assert any("Experienced QA Engineer" in line for line in lines)

def test_generate_docx_adds_footer_if_not_approved():
    text = """
    Jane Smith
    jane@email.com | 555-123-4567

    EXPERIENCE
    Worked at TechCorp
    """
    result = generate_formatted_resume_docx(text, is_user_approved=False)
    doc = Document(result)
    footer = doc.sections[0].footer.paragraphs[0].text
    assert "FindMyDreamJobs" in footer

def test_generate_docx_handles_missing_sections_gracefully():
    text = """
    No obvious section headings here.
    Just plain text describing random things.
    """
    result = generate_formatted_resume_docx(text, is_user_approved=True)
    lines = extract_text_from_docx(result)
    assert any("random things" in line.lower() for line in lines)

def test_generate_docx_formats_bullets_correctly():
    text = """
    Skills
    - Python
    - Testing
    """
    result = generate_formatted_resume_docx(text, is_user_approved=True)
    lines = extract_text_from_docx(result)
    assert any("Python" in line for line in lines)
    assert any("Testing" in line for line in lines)
