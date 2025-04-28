# ✅ File: backend/app/services/resume_formatter.py
# Description: Formats resume text into a Word document with specific styles and sections.
# This module is used to generate a formatted resume in .docx format.
#

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import json
from io import BytesIO
from datetime import datetime
import os
import re
from app.config.skills_config import SECTION_KEYWORDS
from app.config.resume_config import CONF  # style/configuration settingsy
# For hyperlink support
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, url, text, font_size=None):
    """
    Add a hyperlink to a paragraph. Returns the created hyperlink element.
    Optionally specify font_size in points to style the link text.
    """
    # create the relationship id for this hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # build the hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    # create a run for the link
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # blue color
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF'); rPr.append(c)
    # underline style
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    # optional font size
    if font_size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(font_size * 2))); rPr.append(sz)
    new_run.append(rPr)
    # text element
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    # append to paragraph
    paragraph._p.append(hyperlink)
    return hyperlink

def clean_text(raw_text: str) -> str:
    # Collapse multiple blank lines, trim spaces
    cleaned = re.sub(r'\n{2,}', '\n', raw_text.strip())
    return cleaned

def parse_contact_block(lines):
    """
    Extracts the candidate name and a list of contact items (emails, phones, URLs).
    Stops when lines no longer match contact patterns. Returns (name, contact_items, remaining_lines).
    """
    name = ""
    contact_items = []
    # regex patterns for different contact types
    email_re = re.compile(r"[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}")
    phone_re = re.compile(r"(?:\+?\d{1,3}[\s-]?)?(?:\(\d{3}\)|\d{3})[\s\.-]?\d{3}[\s\.-]?\d{4}")
    url_re = re.compile(r"https?://[\w\.-]+(?:/[^\s]*)?")
    if lines:
        name = lines.pop(0).strip()
    # collect subsequent lines that match contact info
    while lines:
        txt = lines[0].strip()
        if email_re.search(txt) or phone_re.search(txt) or url_re.search(txt) or 'linkedin.com' in txt.lower():
            item_line = lines.pop(0).strip()
            # split on pipe or bullet separators into distinct contact segments
            segments = [seg.strip() for seg in re.split(r'\s*(?:\||•)\s*', item_line) if seg.strip()]
            for seg in segments:
                # classify each segment
                if email_re.search(seg):
                    contact_items.append(('email', seg))
                elif url_re.search(seg) or 'linkedin.com' in seg.lower():
                    contact_items.append(('url', seg))
                elif phone_re.search(seg):
                    contact_items.append(('phone', seg))
                else:
                    contact_items.append(('text', seg))
        else:
            break
    return name, contact_items, lines

def is_section_title(line: str) -> bool:
    if not line.strip():
        return False
    words = line.strip().split()
    if len(words) > 5:
        return False
    return any(keyword.lower() in line.lower() for keyword in SECTION_KEYWORDS)

def is_company_location(line: str) -> bool:
    return bool(re.search(r',\s*\w+', line)) and not any(c in line for c in [":", "-", "/"])

def is_job_title_date(line: str) -> bool:
    return bool(re.search(r'\d{2}/\d{4}', line)) or ("present" in line.lower())

def generate_formatted_resume_docx(resume_text: str, is_user_approved: bool) -> BytesIO:
    # Initialize document: use template if specified, else blank
    tpl = CONF.get('template_path')
    if tpl:
        # if relative path, resolve against config directory
        tpl_path = tpl
        if not os.path.isabs(tpl_path):
            tpl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), os.pardir, 'config', tpl)
        try:
            doc = Document(tpl_path)
        except Exception:
            doc = Document()
    else:
        doc = Document()
    output = BytesIO()
    # Apply base font from configuration (if Normal style exists)
    try:
        style = doc.styles['Normal']
        base_font = style.font
        base_font.name = CONF.get('font_name', base_font.name)
    except KeyError:
        # fallback: no Normal style
        pass
    # shorthand access to font sizes
    fs = CONF.get('font_size', {})
    # bullet list style name from config
    bullet_style = CONF.get('styles', {}).get('bullet', 'List Bullet')
    # determine separator image path via config
    img_conf = CONF.get('image_path', None)
    if img_conf:
        # if relative, resolve against this file's folder
        cur = os.path.dirname(os.path.abspath(__file__))
        image_path = img_conf if os.path.isabs(img_conf) else os.path.join(cur, img_conf)
    else:
        image_path = None

    # split into lines and extract name + contact items
    lines = clean_text(resume_text).split('\n')
    name, contact_items, lines = parse_contact_block(lines)

    # Header: Name
    # Header: Name (first page)
    p = doc.add_paragraph()
    run = p.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(fs.get('name', 14))  # configured name font size
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Header: Contact Info (emails, phones & URLs with hyperlinks)
    p = doc.add_paragraph()
    for idx, (ctype, cval) in enumerate(contact_items):
        if ctype == 'email':
            # mailto: hyperlink
            add_hyperlink(p, f"mailto:{cval}", cval, font_size=fs.get('contact', 10))
        elif ctype == 'url':
            # web URL hyperlink
            add_hyperlink(p, cval, cval, font_size=fs.get('contact', 10))
        else:
            run = p.add_run(cval)
            run.font.size = Pt(fs.get('contact', 10))
        # add separator between items
        if idx < len(contact_items) - 1:
            sep = p.add_run(' | ')
            sep.font.size = Pt(fs.get('contact', 10))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set up a different header for pages after the first
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    hdr = section.header
    # add header paragraph for page 2+
    # Header for page 2+: right-aligned header text
    hdr_para = hdr.add_paragraph()
    hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    hdr_run = hdr_para.add_run(CONF.get('header_text', 'Resume of {name}').format(name=name))
    hdr_run.font.size = Pt(fs.get('header', 9))  # configured header font size

    #doc.add_paragraph(" ")
    #doc.add_picture(image_path, width=None, height=None)

    IS_Summary = False
    inside_company_block = False

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        if is_section_title(line):
            inside_company_block = False
            IS_Summary = 'SUMMARY' in line.upper()

            #doc.add_paragraph(" ")
            # insert section separator if available
            if image_path and os.path.exists(image_path):
                doc.add_picture(image_path)

            p = doc.add_paragraph()
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(fs.get('section', 12))  # configured section heading size
            continue

        if is_company_location(line):
            # split company and location, bold only the company
            p = doc.add_paragraph()
            if '|' in line:
                # pipe separator: Company | Location
                parts_pipe = line.split('|', 1)
                company = parts_pipe[0].strip()
                location = parts_pipe[1].strip()
                # bold company
                run = p.add_run(company)
                run.bold = True
                run.font.size = Pt(fs.get('company', 11))  # configured company font size
                # non-bold separator + location
                run2 = p.add_run(' | ' + location)
                run2.bold = False
                run2.font.size = Pt(fs.get('company', 11))  # configured company font size
            else:
                # comma separator: Company, Location
                parts = line.split(',', 1)
                company = parts[0].strip()
                location = parts[1].strip() if len(parts) > 1 else ''
                run = p.add_run(company)
                run.bold = True
                run.font.size = Pt(fs.get('company', 11))  # configured company font size
                if location:
                    run2 = p.add_run(', ' + location)
                    run2.bold = False
                    run2.font.size = Pt(fs.get('company', 11))  # configured company font size
            inside_company_block = True
            continue

        if inside_company_block and is_job_title_date(line):
            # split job title and dates, bold only the title
            p = doc.add_paragraph()
            if '|' in line:
                # pipe separator: Title | Dates
                parts_pipe = line.split('|', 1)
                title = parts_pipe[0].strip()
                dates = parts_pipe[1].strip()
                # bold title
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(fs.get('job', 11))  # configured job title font size
                # non-bold separator + dates
                run2 = p.add_run(' | ' + dates)
                run2.bold = False
                run2.font.size = Pt(fs.get('job', 11))  # configured job title font size
            else:
                # find where the date portion begins
                m = re.search(r'\d{1,2}/\d{4}', line)
                idx = m.start() if m else line.lower().find('present')
                if idx is not None and idx >= 0:
                    # split into title and date, stripping trailing separators
                    title = line[:idx].rstrip(' -–—')
                    dates = line[idx:].strip()
                    # bold title
                    run = p.add_run(title)
                    run.bold = True
                    run.font.size = Pt(fs.get('job', 11))  # configured job title font size
                    # non-bold space + dates
                    run2 = p.add_run(' ' + dates)
                    run2.bold = False
                    run2.font.size = Pt(fs.get('job', 11))  # configured job title font size
                else:
                    # fallback: whole line normal
                    run = p.add_run(line)
                    run.bold = False
                    run.font.size = Pt(11)
            inside_company_block = False
            continue

        if line.startswith("-"):
            # Bulleted list line using configured bullet style
            p = doc.add_paragraph(line[1:].strip(), style=bullet_style)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            if IS_Summary:
                run.font.size = Pt(fs.get('summary', 11))  # configured summary font size
            else:
                run.font.size = Pt(fs.get('bullet', 10))  # configured regular text font size

    # Footer
    # Footer: add approval watermark if needed
    if not is_user_approved:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        run = p.add_run("Generated by FindMyDreamJobs.com – Not yet user approved.")
        run.font.size = Pt(fs.get('footer', 8))  # configured footer font size

    doc.save(output)
    output.seek(0)
    return output

# Reference
resume_formatter_reference = generate_formatted_resume_docx
