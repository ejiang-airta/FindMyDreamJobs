import requests
from bs4 import BeautifulSoup
import schedule
import datetime
import json
import time
import os
from backend.app.config.settings import PROJECT_ROOT
import unicodedata  # Required for Unicode normalization
import logging  # For debugging Unicode issues

# Setup logging for Unicode character debugging
logging.basicConfig(filename='unicode_processing.log', level=logging.DEBUG)

# RapidAPI instance
rapid_api_host = 'jsearch.p.rapidapi.com'
rapid_api_key = os.getenv("RAPIDAPI_KEY")

# Root for storing found jobs: 
logger_dir  = os.path.join(PROJECT_ROOT, "dev_tracking", "jobs")
filename = f"{logger_dir}/app_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

os.makedirs(logger_dir , exist_ok=True)
def fetch_jobs():
    job_postings = []
    API_URL = f"https://{rapid_api_host}/search"    #jsearch API URL

    
    # jsearch API
    headers = {
        'x-rapidapi-host': rapid_api_host.encode("ascii", "ignore").decode(),
        'x-rapidapi-key': rapid_api_key.encode("ascii", "ignore").decode()
        }
    # API parameters
    params = {
            "query":"director engineering jobs in Vancouver, BC",
            "page":"1",
            "num_pages":"9",
            "country":"ca",
            "date_posted":"all",
            "work_from_home":"true"
    }
    
    # Make the API request:
    # response = requests.get(API_URL, headers=headers,params=params)
    # data = response.json()
    # Mock up the API response:
    with open('jobs_list_1.json', 'r') as file:
        data = json.load(file)
    
    #job_postings.extend(data['results'])
    job_postings = data.get("data")

    # Save raw response for debugging
    with open('job_output.json', 'w', encoding='utf-8') as file1:
        json.dump(data, file1, indent=4, ensure_ascii=False) # Pretty-print
        print("✅ Job information have been written to job_output.json")

    return job_postings

def display_jobs(jobs):
    # Display jobs in a user-friendly format
    print("*****Printing Job posting*****")
    iStart = 1

    for job in jobs:        
        print(f"***Fond job Number: {iStart}***")
        print(f"**Title: {job['job_title']}**")
        print(f"**Location: {job.get('job_location', 'N/A')}**")
        print(f"**Company: {job.get('employer_name', 'N/A')}**")
        print("-"*30)
        print(f"***Job Description: {job.get('job_description', 'N/A')}***")
        print("-"*30)
        print(f"**Posted at: {job.get('job_posted_at_datetime_utc', 'N/A')}**")
        print(f"**Salary: {job.get('job_salary', 'N/A')}**")
        print(f"**Min Salary: {job.get('job_min_salary', 'N/A')}**")
        print(f"**Max Salary: {job.get('job_max_salary', 'N/A')}**")
        print(f"**Apply at: <a href='{job.get('job_google_link', '#')}' target='_blank'>Click here to apply</a>**")
        print("-"*60)

        iStart += 1    
    

# Fetch jobs
jobs = fetch_jobs()

# Sort jobs by post date descending
def parse_date(dt_str):
    """Parse an ISO datetime string into a naive datetime in UTC for sorting."""
    try:
        if not dt_str:
            return datetime.datetime.min
       
       # Parse RFC3339-like with 'Z' suffix as UTC
        dt = datetime.datetime.fromisoformat(dt_str.replace('Z', '+00:00'))
        # Convert to UTC and drop tzinfo to produce naive datetime
        if dt.tzinfo is not None:
            dt = dt.astimezone(datetime.timezone.utc).replace(tzinfo=None)
        return dt
    except Exception:
        # Fallback minimal datetime
        return datetime.datetime.min
    
# Sort jobs by post date descending
jobs = sorted(
    jobs,
    key=lambda j: parse_date(j.get('job_posted_at_datetime_utc', '') or ''),
    reverse=True
)

# Display sorted jobs
display_jobs(jobs)

# Save jobs to HTML file with clickable links
def save_jobs_to_html(jobs, filename='job_output.html'):
    """Generate an HTML file from job postings with clickable links."""
    html_lines = [
        '<!DOCTYPE html>',
        '<html lang="en">',
        '<head><meta charset="UTF-8"><title>Job Postings</title></head>',
        '<body>',
        '<h1>Job Postings</h1>'
    ]
    for idx, job in enumerate(jobs, start=1):
        title = job.get('job_title', 'N/A')
        company = job.get('employer_name', 'N/A')
        location = job.get('job_location', 'N/A')
        posted = job.get('job_posted_at_datetime_utc', 'N/A')
        salary = job.get('job_salary', 'N/A')
        min_salary = job.get('job_min_salary', 'N/A')
        max_salary = job.get('job_max_salary', 'N/A')
        description = job.get('job_description', '')
        link = job.get('job_google_link', '#')

        # Job title header
        html_lines.append(f'<h2>Job {idx}: {title}</h2>')
        # Bullet list of metadata
        html_lines.append('<ul>')
        for label, value in [
            ('Company', company),
            ('Location', location),
            ('Posted at', posted),
            ('Salary', salary),
            ('Min Salary', min_salary),
            ('Max Salary', max_salary)
        ]:
            html_lines.append(f'<li><strong>{label}:</strong> {value}</li>')
        html_lines.append('</ul>')
        # Description with paragraphs
        html_lines.append('<h3>Description:</h3>')
        for line in description.split('\n'):
            if line.strip():
                html_lines.append(f'<p>{line}</p>')
        # Apply link
        html_lines.append(f'<p><strong>Apply at:</strong> <a href="{link}" target="_blank">Click here to apply</a></p>')
        # Separator
        html_lines.append('<hr>')
    html_lines.extend(['</body>', '</html>'])
    with open(filename, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_lines))
    print(f"✅ Job information saved to {filename}")

save_jobs_to_html(jobs)

# Unicode-safe DOCX export
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
except ImportError:
    print("⚠️  python-docx not installed. Skipping DOCX export. Install with `pip install python-docx` to enable this feature.")
else:
    def normalize_string(text):
        """
        Unicode-safe string normalization with fallback handling
        Fixes issues with literal unicode escapes (e.g., \uD83D\uDC76) and invalid surrogate pairs
        """
        if not isinstance(text, str):
            return str(text)

        # Decode literal unicode escape sequences (e.g., \uXXXX or \UXXXXXXXX) into actual characters
        if '\\u' in text or '\\U' in text:
            try:
                # operate on utf-8 bytes to use unicode_escape codec
                text = text.encode('utf-8').decode('unicode_escape')
            except Exception:
                logging.debug("normalize_string: failed to decode unicode escapes")

        # Normalize to NFKC form to handle compatibility characters
        text = unicodedata.normalize('NFKC', text)

        # Special character replacements needed for DOCX
        replacements = {
            '\u200b': '',  # Zero-width space
            '\ufeff': '',   # Byte order mark
            '\u202a': '',   # Left-to-right embedding
            '\u202c': '',   # Pop directional formatting
        }
        # Apply replacements
        for old, new in replacements.items():
            text = text.replace(old, new)

        # Additional check to handle invalid surrogate pairs
        try:
            # Try to encode using utf-16 to catch invalid surrogates
            return text.encode('utf-16', 'surrogatepass').decode('utf-16')
        except UnicodeDecodeError:
            logging.warning("Failed to handle surrogate pairs in text")
            # Fallback for problematic cases
            return text.encode('utf-8', 'replace').decode('utf-8', 'replace')
            
    def find_invalid_chars(text):
        """For debugging: identify problematic Unicode characters"""
        if not isinstance(text, str):
            return []
        invalid_chars = []
        for char in set(text):
            try:
                char.encode('utf-16', 'surrogatepass')
            except UnicodeEncodeError:
                invalid_chars.append(repr(char))
        return invalid_chars    
    
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a python-docx paragraph."""
        processed_text = normalize_string(text)
        processed_url = normalize_string(url)

        part = paragraph.part
        r_id = part.relate_to(processed_url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF')
        u = OxmlElement('w:u');    u.set(qn('w:val'), 'single')
        rPr.append(c); rPr.append(u)
        new_run.append(rPr)
        text_el = OxmlElement('w:t'); text_el.text = processed_text
        new_run.append(text_el)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def save_jobs_to_docx(jobs, filename):
        """Generate a DOCX file from job postings with formatted headers and bullet points and with Unicode-safe formatting."""
        print(f"Docx export file name: {filename}")
        doc = Document()
        
        # Set compatibility mode for better Unicode handling
        doc._element.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}compatibilityMode', 
            '15'
        )
        doc.add_heading('Job Postings', level=1)
        total = len(jobs)
        for idx, job in enumerate(jobs, start=1):
        # Process job fields with Unicode normalization
            title = normalize_string(job.get('job_title', 'N/A'))
            company = normalize_string(job.get('employer_name', 'N/A'))
            location = normalize_string(job.get('job_location', 'N/A'))
            posted = normalize_string(job.get('job_posted_at_datetime_utc', 'N/A'))
            salary = normalize_string(job.get('job_salary', 'N/A'))
            min_salary = normalize_string(job.get('job_min_salary', 'N/A'))
            max_salary = normalize_string(job.get('job_max_salary', 'N/A'))
            description = job.get('job_description', '')
            link = normalize_string(job.get('job_google_link', '#'))
            
            # Log invalid characters for debugging
            for field_name, field_value in [
                ('title', title),
                ('company', company),
                ('link', link)
            ]:
                invalid_chars = find_invalid_chars(field_value)
                if invalid_chars:
                    logging.debug(f"Invalid chars in job {idx} {field_name}: {', '.join(invalid_chars)}")
            
            # Add heading with normalized text
            doc.add_heading(f'Job {idx}: {title}', level=2)
            for label, value in [
                ('Company', company),
                ('Location', location),
                ('Posted at', posted),
                ('Salary', salary),
                ('Min Salary', min_salary),
                ('Max Salary', max_salary)
            ]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{label}: {value}')

            # Process description with line-by-line normalization
            if description:
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run('Description:')
                desc_run.bold = True
                
                for line in description.split('\n'):
                    if line.strip():
                        processed_line = normalize_string(line)
                        doc.add_paragraph(processed_line)

            # Add hyperlink with normalized text
            link_para = doc.add_paragraph()
            link_para.add_run('Apply at: ')
            add_hyperlink(link_para, 'Click here to apply', link)

            # Add page break between jobs except for last one
            if idx < total:
                doc.add_page_break()

        doc.save(filename)
        print(f"✅ Job information saved to {filename}")

    save_jobs_to_docx(jobs,filename=filename)
# def main():
#     schedule.every(1).day.at("11:36").do(fetch_jobs)  # Run daily at 8 AM
#     while True:
#         schedule.run_pending()
#         time.sleep(60)

# if __name__ == "__main__":
#     main()